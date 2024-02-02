import streamlit as st
import PyPDF2
import docx
from openai import OpenAI
import time
import os
from llama_index import Document, VectorStoreIndex, SimpleDirectoryReader
from llama_index.node_parser import SentenceSplitter

client = OpenAI(api_key='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')
OPENAI_API_KEY = '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'
os.environ["OPENAI_API_KEY"] = OPENAI_API_KEY


def initial_message():
    st.text("Assistant: Hello! I'm here to assist you.")


def load_files():
    uploaded_files = st.file_uploader("Pick one or more files to test", type=["pdf", "txt", "docx"], accept_multiple_files=True)

    if uploaded_files:
        st.success("Files uploaded successfully!")

        # Determine the file types and display content accordingly
        file_texts = []
        for uploaded_file in uploaded_files:
            file_extension = uploaded_file.name.split('.')[-1].lower()
            st.text(f"Acknowledgement: {file_extension.capitalize()} file uploaded.")
            file_texts.append((uploaded_file, file_extension))

        return file_texts
    else:
        return []


def chat_with_openai(user_message, conversation_history_container, document_texts):
    # Create a text_input widget for user input
    user_input = st.text_input("User:", value=user_message)

    # Use a unique key for the button to avoid DuplicateWidgetID error
    button_key = "send_button_" + str(hash(user_message))
    if st.button("Send", key=button_key):
        messages = []
        system_msg = "Your system message here"  # Provide an appropriate system message
        messages.append({"role": "system", "content": system_msg})
        messages.append({"role": "user", "content": user_input})

        # Include document content in the conversation history
        for document_text in document_texts:
            messages.append({"role": "user", "content": document_text})  # Use 'user' role for document content

        try:
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=messages
            )
            reply = response.choices[0].message.content
            messages.append({"role": "assistant", "content": reply})

            # Display user input, document content, and assistant's reply in conversation history
            current_content = conversation_history_container.text_area("Conversation History:", value="", height=300,
                                                                      max_chars=1000)
            conversation_history_container.text_area("Conversation History:",
                                                    value=f"{current_content}\nUser: {user_input}\nAssistant: {reply}",
                                                    height=300, max_chars=1000)
        except Exception as e:
            st.error("An error occurred: " + str(e))
            print("An error occurred: ", e)


def create_index_from_uploaded_files(uploaded_files):
    documents = []
    file_texts = []

    for uploaded_file, file_extension in uploaded_files:
        if file_extension == "txt":
            # Assuming a plain text file
            text = uploaded_file.getvalue().decode("utf-8")
        elif file_extension == "pdf":
            # Assuming a PDF file
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                text += pdf_reader.pages[page_num].extract_text()
        elif file_extension == "docx":
            # Assuming a DOCX file
            doc = docx.Document(uploaded_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        else:
            st.error(f"Unsupported file type: {file_extension}")
            continue

        documents.append(Document(text=text, id_=uploaded_file.name))
        file_texts.append(text)

    # Create and return the index
    parser = SentenceSplitter()
    nodes = parser.get_nodes_from_documents(documents)
    return VectorStoreIndex(nodes), file_texts


# This function is used to query the index and retrieve relevant documents
def query_index(query, index):
    if index is not None:
        return index.search(query)
    else:
        return []


def main():
    st.title('File Upload and Chat Demo')

    # Initial message from the assistant
    initial_message()

    # Allow users to upload multiple files (PDF, TXT, DOCX)
    uploaded_files = load_files()

    # Use the uploaded files to create the index
    index, document_texts = create_index_from_uploaded_files(uploaded_files)

    # Chat with the assistant
    conversation_history_container = st.empty()

    # Always display the text input box
    chat_with_openai("", conversation_history_container, document_texts)


if __name__ == '__main__':
    main()
